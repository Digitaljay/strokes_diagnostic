from docx import Document
import datetime
class Resolution():
    def __init__(self, name="Григорий",
                       age=70,
                       recs=["-"],
                       patology="Ишемический инсульт",
                       neurological_deficit=1,              # [1, 2, 3, 4]
                       conscious_level=15,                  # [15, 14, 12, 9, 7, 5, 3]
                       time_passed=0,                       # [int]
                       hematoma_volume=0,                   # [int]
                       is_injury=None,                      # [True/False/None]
                       has_stroke_symptoms=True,            # [True/False]
                       chronic=None,                        # [[*]/None]
                       temporary_contraindications=None):   # [[*]/None]

        self.patology=patology
        self.neurological_deficit=neurological_deficit
        self.conscious_level=conscious_level
        self.time_passed = time_passed
        self.hematoma_volume=hematoma_volume
        self.is_injury=is_injury
        self.has_stroke_symptoms=has_stroke_symptoms
        self.chronic=chronic
        self.temporary_contraindications=temporary_contraindications
        self.name=name
        self.age=age
        self.recs=recs

        self.document = Document()
        self.document.add_heading('Заключение', 0)

        p1=self.document.add_paragraph('')
        p1.add_run("Пациент: ").bold=True
        p1.add_run(self.name)

        p2=self.document.add_paragraph('')
        p2.add_run("Дата пождения: ").bold=True

        p3=self.document.add_paragraph('')
        p3.add_run("Возраст: ").bold=True
        p3.add_run(str(self.age)+" лет")

        p4=self.document.add_paragraph('')
        p4.add_run("Дата консультации: ").bold=True
        now = datetime.datetime.now()
        self.date=str(now.strftime("%d-%m-%Y %H:%M")).split()[0]
        p4.add_run(self.date)

        p5=self.document.add_paragraph('')
        p5.add_run("Анамнез: ").bold=True

        self.document.add_paragraph("Сопутствующие патологи: ")
        for pat in self.chronic:
            self.document.add_paragraph(pat, style='List Bullet')

        self.document.add_paragraph("Противопоказания: ")
        for pat in self.temporary_contraindications:
            self.document.add_paragraph(pat, style='List Bullet')

        p6=self.document.add_paragraph('')
        p6.add_run("Данные обследования: ").bold=True
        self.document.add_paragraph("Уровень сознания: "+ str(self.conscious_level))
        self.document.add_paragraph("Неврологический дефицит: "+str(self.neurological_deficit))
        if self.has_stroke_symptoms:
            self.document.add_paragraph("Присутствуют симптомы инсульта")
        else:
            self.document.add_paragraph("Симптомы инсульта отсутствуют")
        if self.is_injury:
            self.document.add_paragraph("Была получена травма")

        p7=self.document.add_paragraph('')
        p7.add_run("Диагноз: ").bold=True
        p7.add_run(self.patology)

        p8=self.document.add_paragraph('')
        p8.add_run("Рекомендовано: ").bold=True
        for rec in recs:
            self.document.add_paragraph(rec, style='List Bullet')


        self.document.save('medical_resolution.docx')

